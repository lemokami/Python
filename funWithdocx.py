#Modifying creating and handling word documents with docx library
import docx #importing docx

doc = docx.Document() #creates a document if path is not mentioned in the parenthesis else creates a new instance


doc.add_paragraph("HEllo I'm the heading")#adds a new Paragraph to the file


para = doc.paragraphs[0] #selecting the first pargraph
para.add_run("\nHello from the content.\nHow are you there little visitor") #adding the run to the selected paragraph

para.runs[0].bold = True #makes the first run i.e the first paragraph(heading) bold 
para.runs[1].italic = True #makes the first run i.e the heading bold

for paraG in doc.paragraphs: #running a loop through the the paragraphs in the file
    print(paraG.text) #printing the text with .text since just printing with paraG iterable prints some weird stuff ;)

doc.save("Files\\Docx\\demo.docx")
#since we are making an instance an changing it here it is not saved in the local directory until we do .save operation like in other software so we did it here
print("\n\nPrinting finished before this line\n\nFile saved in Docx Folder in the Files Folder(Files\\Docx or Files/Docx)")